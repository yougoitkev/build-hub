from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from PIL import Image

def prepare_image(image_path):
    """Re-saves the image using Pillow to ensure it's a standard PNG."""
    if not os.path.exists(image_path):
        print(f"File not found: {image_path}")
        return None
    try:
        img = Image.open(image_path)
        tmp_path = image_path + ".v4pr.png"
        img.save(tmp_path, "PNG")
        return tmp_path
    except Exception as e:
        print(f"Error processing {image_path}: {e}")
        return None

def add_styled_section(doc, title, level=1):
    heading = doc.add_heading(title, level=level)
    return heading

def add_operational_note(doc, intent, before=None, after=None):
    p = doc.add_paragraph()
    p.add_run("👉 User Intent: ").bold = True
    p.add_run(intent)
    
    if before:
        p = doc.add_paragraph()
        p.add_run("⬅️ Before this step: ").italic = True
        p.add_run(before)
    
    if after:
        p = doc.add_paragraph()
        p.add_run("➡️ After this step: ").italic = True
        p.add_run(after)

def add_decision_points(doc, scenarios):
    doc.add_heading("Decision Points & Troubleshooting", level=3)
    for scenario, action in scenarios.items():
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"Scenario: {scenario}\n").bold = True
        p.add_run(f"Action: {action}").italic = True

def create_manual():
    doc = Document()
    brain_dir = r"C:\Users\349125\.gemini\antigravity\brain\333ecd82-2881-4b60-8186-e6dcb9ab2723"

    # Screenshot Paths
    img_login = prepare_image(os.path.join(brain_dir, "media__1773845937331.png"))
    img_performance_sup = prepare_image(os.path.join(brain_dir, "media__1773845937496.png"))
    img_add_trainer = prepare_image(os.path.join(brain_dir, "media__1773845937643.png"))
    img_calendar_sup = prepare_image(os.path.join(brain_dir, "media__1773845937781.png"))
    img_dashboard_sup = prepare_image(os.path.join(brain_dir, "media__1773845937782.png"))
    img_learners_sup = prepare_image(os.path.join(brain_dir, "media__1773847602762.png"))
    img_audit = prepare_image(os.path.join(brain_dir, "media__1773847602922.png"))

    img_dashboard_tra = prepare_image(os.path.join(brain_dir, "media__1773847603088.png"))
    img_learners_tra = prepare_image(os.path.join(brain_dir, "media__1773847603260.png"))
    img_new_student = prepare_image(os.path.join(brain_dir, "media__1773848065120.png"))
    img_calendar_tra = prepare_image(os.path.join(brain_dir, "media__1773848065282.png"))
    img_attendance_tra = prepare_image(os.path.join(brain_dir, "media__1773848065438.png"))
    img_observations_tra = prepare_image(os.path.join(brain_dir, "media__1773848065455.png"))
    img_performance_tra = prepare_image(os.path.join(brain_dir, "media__1773848065496.png"))

    # 1. Cover Page
    doc.add_heading('Training Management System (TMS)', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Operational User Guide — Real-World Workflows').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n' * 3)
    doc.add_paragraph('Transforming Training Operations from Planning to Performance').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n' * 2)
    doc.add_paragraph('Version 2.0 (Enhanced)').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Date: March 20, 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 2. Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc = [
        "1. Introduction", 
        "2. How the System Works Together (Cross-Module Map)",
        "3. Login & Authentication", 
        "4. Supervisor Full Flow: Monitoring Operations", 
        "5. Trainer Full Flow: Running a Training Day",
        "6. Detailed Module Guide: Supervisor",
        "7. Detailed Module Guide: Trainer",
        "8. Operational Troubleshooting"
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph("Welcome to the TMS Operational Guide. This document is designed for first-time users to master the system through real-world scenarios. Unlike a technical manual, this guide focuses on 'Why' and 'How' you perform your daily tasks.")

    # 2. How the System Works Together
    doc.add_heading('2. How the System Works Together', level=1)
    doc.add_paragraph("The TMS is a closed-loop ecosystem where data flows seamlessly between modules:")
    flow_steps = [
        "Program Creation: The starting point where templates define session schedules.",
        "Calendar Allocation: Programs automatically populate the Calendar for Trainers.",
        "Daily Execution: Trainers use the Calendar to launch Attendance and Observations.",
        "Performance Visibility: Attendance and Observation data aggregate into Performance Gantt charts for both Trainers and Supervisors."
    ]
    for step in flow_steps:
        doc.add_paragraph(step, style='List Number')

    # 3. Login
    doc.add_heading('3. Login & Authentication', level=1)
    add_operational_note(doc, "Securely access your personalized workspace based on your role (Supervisor or Trainer).", None, "Dashboard access.")
    if img_login:
        doc.add_picture(img_login, width=Inches(5.5))
    doc.add_paragraph('Figure 1: Secure Login Interface').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 4. Supervisor Full Flow
    doc.add_heading('4. Supervisor Full Flow: Managing Trainers & Monitoring Operations', level=1)
    doc.add_paragraph("Goal: Ensure all training programs are staffed, scheduled, and progressing according to plan.")
    steps = [
        "Login: Access the administrative command center.",
        "Dashboard Review: Check the 'System Summary' to identify any immediate red flags or pending actions.",
        "Trainer Creation: If a new trainer joins, go to the 'Trainer Form' and create their profile. This permits them to be assigned to programs.",
        "Verification: Confirm the new trainer appears in the 'Trainer Directory'.",
        "Resource Monitoring: Open the 'Calendar' to see global class distribution. Ensure no overlaps and that all trainers have balanced workloads.",
        "In-Depth Analysis: Navigate to 'Performance' to review the Gantt timeline. This shows the long-term roadmap for all active cohorts.",
        "Accountability: Check the 'Audit Trail' to verify that attendance and observations are being recorded on time by the training team."
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Bullet')

    # 5. Trainer Full Flow
    doc.add_heading('5. Trainer Full Flow: Running a Training Program (End-to-End)', level=1)
    doc.add_paragraph("Goal: Execute a training program effectively, from first-hour attendance to end-of-day observations.")
    
    doc.add_heading("Scenario: A Standard Day (9 AM - 5 PM)", level=2)
    daily_steps = [
        "09:00 AM - Shift Start: Log in and land on your 'Trainer Dashboard'. Review 'Today’s Schedule' in the sidebar.",
        "09:15 AM - Prep: Go to 'My Calendar'. Select today's session to confirm student enrollment and session objectives.",
        "09:30 AM - Enrollment Check: Navigate to the 'Students' page. If a new student joins late, use the 'New Student' form to add them immediately.",
        "10:00 AM - Attendance: Open the 'Attendance Matrix'. Select today’s date and training program. Record hours for each student (e.g., 8 hours for full attendance, 0 for NCNS). Save your changes.",
        "01:00 PM - Progress Tracking: Periodically check the 'Performance' page to see your program's progress against the original template.",
        "04:30 PM - Observations: Before the day ends, open the 'Observation' page. Rate students on Behavior, Soft Skills, and Participation. Add qualitative notes for follow-up. Save for system visibility.",
        "05:00 PM - Wrap Up: Verify all entries are saved. Your data is now visible to the Supervisor for their daily reporting."
    ]
    for s in daily_steps:
        doc.add_paragraph(s, style='List Bullet')

    # 6. Detailed Module Guide: Supervisor
    doc.add_heading('6. Detailed Module Guide: Supervisor', level=1)
    
    doc.add_heading('6.1 Dashboard (Command Center)', level=2)
    add_operational_note(doc, "Get an instant pulse on training health without digging through menus.", "Daily Login", "Resource deep-dive.")
    if img_dashboard_sup:
        doc.add_picture(img_dashboard_sup, width=Inches(5.5))
    doc.add_paragraph("Usage: Review the 'Summary' widgets to see how many trainers are active and how many students are currently in the system.")

    doc.add_heading('6.2 Trainer Management', level=2)
    add_operational_note(doc, "Maintain the roster of available trainers and their specializations.", "New Hire Request", "Program Assignment.")
    if img_add_trainer:
        doc.add_picture(img_add_trainer, width=Inches(5.5))
    add_decision_points(doc, {
        "Trainer Email missing?": "The system requires a valid email to link notifications.",
        "Trainer status changed?": "Flip the toggle to 'Inactive' to prevent further program assignments."
    })

    doc.add_heading('6.3 Global Calendar View', level=2)
    add_operational_note(doc, "Visual heatmap of all sessions. Avoids scheduling conflicts and site congestion.", "Strategic Planning", "Operational Review.")
    if img_calendar_sup:
        doc.add_picture(img_calendar_sup, width=Inches(5.5))
    doc.add_paragraph("The legend color-codes sessions by Primary Trainer, allowing for instant recognition of who is leading which class.")

    # 7. Detailed Module Guide: Trainer
    doc.add_heading('7. Detailed Module Guide: Trainer', level=1)
    
    doc.add_heading('7.1 Learner Management & Enrollment', level=2)
    add_operational_note(doc, "Maintain accurate records of who is in your class.", "Program Launch", "Attendance Tracking.")
    if img_learners_tra:
        doc.add_picture(img_learners_tra, width=Inches(5.5))
    if img_new_student:
        doc.add_picture(img_new_student, width=Inches(5.5))
    add_decision_points(doc, {
        "Student not in list?": "Verify they are enrolled in the correct training ID via the Student Directory.",
        "Multiple Enrollments?": "Check the 'Active Only' filter to see the current program."
    })

    doc.add_heading('7.2 Attendance Matrix', level=2)
    add_operational_note(doc, "Crucial for payroll and compliance. Marks the physical presence of students.", "Class Start", "Observation Rating.")
    if img_attendance_tra:
        doc.add_picture(img_attendance_tra, width=Inches(5.5))
    add_decision_points(doc, {
        "Wrong date selected?": "Use the date picker in the matrix header to switch to the correct session.",
        "Bulk update needed?": "Select multiple cells to apply the same hours/status simultaneously."
    })

    doc.add_heading('7.3 Daily Student Observations', level=2)
    add_operational_note(doc, "Qualitative scoring of performance. Bridges the gap between 'Being Present' and 'Being Proficient'.", "Attendance Marked", "Performance Review.")
    if img_observations_tra:
        doc.add_picture(img_observations_tra, width=Inches(5.5))
    doc.add_paragraph("Example: A student might be Present (8 hrs) but receive a 'Needs Improvement' in Performance due to low engagement.")

    doc.add_heading('7.4 My Performance Timeline', level=2)
    add_operational_note(doc, "Visualizing your own delivery speed and milestones.", "Daily Wrap up", "Weekly Review.")
    if img_performance_tra:
        doc.add_picture(img_performance_tra, width=Inches(5.5))

    # 8. Troubleshooting
    doc.add_heading('8. Operational Troubleshooting', level=1)
    troubleshoot = {
        "No students showing in attendance?": "Verify the Program Start Date. If the program hasn't started yet, the attendance matrix will remain empty for those dates.",
        "Calendar appears empty?": "Ensure you are filtering for your own Trainer ID or that a Training Program has been properly 'Scheduled' in the system.",
        "Performance data not updating?": "Data typically refreshes after saving Attendance and Observation forms. Ensure you clicked 'Save' in the respective modules.",
        "Wrong session type in Calendar?": "Verify the 'Template' used during program creation. Templates define the sequence and type of sessions."
    }
    for issue, solution in troubleshoot.items():
        p = doc.add_paragraph()
        p.add_run(f"Issue: {issue}\n").bold = True
        p.add_run(f"Solution: {solution}").italic = True

    # Save
    save_path = os.path.join(r"C:\Wokspace\Mar'26\TMS\trainer-hub-main", "TMS_User_Manual_Final.docx")
    doc.save(save_path)
    print(f"Updated manual saved to {save_path}")

if __name__ == "__main__":
    create_manual()
